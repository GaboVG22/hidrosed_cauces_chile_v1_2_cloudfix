"""Generación de informe técnico Word."""
from __future__ import annotations
from pathlib import Path
from typing import Dict
import pandas as pd


def generate_word_report(output_path: str | Path, results: Dict[str, pd.DataFrame], warnings: list[str] | None = None) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        txt_path = output_path.with_suffix(".txt")
        txt_path.write_text("python-docx no disponible.\n" + "\n".join(warnings or []), encoding="utf-8")
        return txt_path
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    doc.add_heading("Informe técnico preliminar - HidroSed Cauces Chile", 0)
    doc.add_paragraph("Herramienta de apoyo técnico para hidrología, hidráulica 1D, sedimentos, socavación, depositación e inundación preliminar en planta.")
    doc.add_heading("Supuestos y advertencias", level=1)
    for w in warnings or []:
        doc.add_paragraph(str(w), style=None)
    doc.add_paragraph("Los resultados son preliminares y deben ser revisados por especialista. No reemplazan modelación HEC-RAS 1D/2D ni levantamiento topográfico de detalle.")
    for name, df in results.items():
        doc.add_heading(str(name), level=1)
        if df is None or df.empty:
            doc.add_paragraph("Sin resultados.")
            continue
        show = df.head(15)
        table = doc.add_table(rows=1, cols=len(show.columns))
        table.style = "Table Grid"
        for j, col in enumerate(show.columns):
            table.rows[0].cells[j].text = str(col)
        for _, row in show.iterrows():
            cells = table.add_row().cells
            for j, col in enumerate(show.columns):
                val = row[col]
                if isinstance(val, float):
                    cells[j].text = f"{val:.3f}"
                else:
                    cells[j].text = str(val)
    doc.save(output_path)
    return output_path
